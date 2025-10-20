# utils.py
import re
import logging
from uuid import UUID
from datetime import datetime
from typing import Any, Dict, List
import config
import numpy as np
# --- NEW IMPORTS for PDF parsing ---
import fitz  # PyMuPDF
import pypdfium2 as pdfium
from PIL import Image
# --- END NEW IMPORTS ---

LOG = logging.getLogger("uvicorn.error")

# =============================================================================
# 1. Data Cleaning and Sanitization Utilities
# =============================================================================

def sanitize_for_json(obj: Any) -> Any:
    """
    Recursively converts common non-serializable types into JSON-serializable types.
    Handles numpy arrays/scalars, UUIDs, and datetimes.
    Also truncates long text fields to keep the final JSON response manageable.
    """
    # Primitives
    if obj is None or isinstance(obj, (str, bool, int, float)):
        return obj

    # Numpy arrays and scalars
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    if isinstance(obj, np.generic):
        return obj.item()

    # UUID and datetime objects
    if isinstance(obj, UUID):
        return str(obj)
    if isinstance(obj, datetime):
        return obj.isoformat()

    # Dictionaries: recurse and handle special keys
    if isinstance(obj, dict):
        out = {}
        for k, v in obj.items():
            # Skip huge binary fields like embeddings
            if k == "embedding":
                continue
            # Truncate potentially long text fields to avoid bloating the response
            if k == "orig_text" or (k == "metadata" and isinstance(v, dict) and "text" in v):
                if k == "orig_text":
                    out[k] = (v or "")[:config.MAX_TEXT_CHARS]
                else: # It's metadata with a 'text' key
                    md = {mdk: (mdv or "")[:config.MAX_TEXT_CHARS] if mdk == "text" else sanitize_for_json(mdv) for mdk, mdv in v.items()}
                    out[k] = md
                continue
            out[k] = sanitize_for_json(v)
        return out

    # Lists, tuples, sets: recurse
    if isinstance(obj, (list, tuple, set)):
        return [sanitize_for_json(x) for x in obj]

    # Fallback for other objects
    if hasattr(obj, "__dict__"):
        return sanitize_for_json(vars(obj))
    
    return str(obj)


def _post_process_and_clean_questions(questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Cleans a list of parsed questions to fix common LLM generation errors,
    such as duplicate internal choices and badly formatted case studies.
    """
    if not questions:
        return []

    cleaned_questions = []
    for question in questions:
        # FIX 1: Collapse duplicate internal choices
        if question.get("is_choice") and isinstance(question.get("q_text"), list) and len(question["q_text"]) > 1:
            choice_texts = [choice.get("q_text", "").strip() for choice in question["q_text"]]
            if len(set(choice_texts)) == 1:
                LOG.warning(f"Fixing duplicate internal choice for q_id: {question.get('q_id')}")
                question["q_text"] = choice_texts[0]
                del question["is_choice"]
                question["_cleaning_note"] = "Collapsed duplicate internal choices."

        # FIX 2: Restructure badly formatted Case Studies
        if "case" in question.get("type", "").lower() and isinstance(question.get("q_text"), list):
            all_sub_texts = [item.get("q_text", "") for item in question.get("q_text", [])]
            if all_sub_texts:
                first_full_text = all_sub_texts[0]
                match = re.search(r"\(\s*[ivx]+\s*\)", first_full_text)
                if match:
                    passage = first_full_text[:match.start()].strip()
                    cleaned_sub_questions = [text.replace(passage, "").strip() for text in all_sub_texts if text.replace(passage, "").strip()]
                    if passage and cleaned_sub_questions:
                        LOG.warning(f"Restructuring case study for q_id: {question.get('q_id')}")
                        question["q_text"] = {"passage": passage, "questions": cleaned_sub_questions}
                        question["_cleaning_note"] = "Restructured case study to proper passage/questions format."

        cleaned_questions.append(question)
    return cleaned_questions


# =============================================================================
# 2. Text and String Manipulation Utilities
# =============================================================================

def standardize_subject_name(subject: str) -> str:
    """
    Maps a potentially messy user-provided subject string to a clean,
    standardized value for metadata filtering (e.g., in Pinecone).
    """
    s_lower = subject.strip().lower()
    subject_map = {
        "maths": "Maths",
        "mathematics": "Maths",
        "math": "Maths",
        "physics": "Physics",
        "science": "Science",
        "chemistry": "Chemistry",
        "biology": "Biology",
    }
    return subject_map.get(s_lower, subject.strip().title())


def generate_tokens(text: str, chunk_size: int = 500) -> List[str]:
    """Splits a string of text into a list of smaller chunks."""
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]


# =============================================================================
# 3. Mathematical and Vector Utilities (from LLM Researcher)
# =============================================================================

def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    """Calculates the cosine similarity between two numpy vectors."""
    a = np.array(a, dtype=float)
    



    
# =============================================================================
# 4. Document Parsing Utilities
# =============================================================================

def parse_pdf(file_content: bytes = None, filename: str = None, file_path: str = None) -> Dict[str, Any]:
    """
    Parses a PDF file provided either as bytes or as a file path. Extracts text if
    available; otherwise converts pages to images. Accepts either `file_content`
    (bytes) or `file_path` (str). If both are provided, `file_path` takes precedence.
    """
    try:
        # Prefer file_path when available to avoid loading bytes into memory
        if file_path:
            pdf_document = fitz.open(file_path)
        else:
            pdf_document = fitz.open(stream=file_content, filetype="pdf")

        text = ""
        for page in pdf_document:
            text += page.get_text()

        if text.strip():
            LOG.info(f"Successfully extracted text from PDF '{filename or file_path}'.")
            return {"text": text}
        else:
            LOG.info(f"No text found in PDF '{filename or file_path}', processing as image-based.")
            images = []
            # pdfium can accept a filename or bytes; prefer filepath to avoid extra memory
            if file_path:
                pdf_pages = pdfium.PdfDocument(file_path)
            else:
                pdf_pages = pdfium.PdfDocument(file_content)

            for i in range(len(pdf_pages)):
                page = pdf_pages.get_page(i)
                bitmap = page.render(scale=2)  # Using a good scale for clarity
                pil_image = bitmap.to_pil()
                images.append(pil_image)
            return {"images": images}
    except Exception as e:
        LOG.error(f"Error parsing PDF '{filename or file_path}': {e}")
        # Re-raise the exception to be handled by the service layer
        raise ValueError(f"Error parsing PDF file: {filename or file_path}") from e




# In utils.py, add these new imports at the top
import os
import io
from docx import Document

# In utils.py, add this new function, likely in the "Document Parsing Utilities" section

def parse_document_to_text(file_content: bytes = None, filename: str = None, file_path: str = None) -> str:
    """
    Parses an uploaded document (PDF, DOCX, TXT) and extracts the raw text.
    """
    text = ""

    LOG.info(f"Parsing document '{filename or file_path}'.")

    try:
        # Determine extension from filename or path
        if filename:
            file_extension = os.path.splitext(filename)[1].lower()
        elif file_path:
            file_extension = os.path.splitext(file_path)[1].lower()
        else:
            raise ValueError("No filename or file_path provided to parse_document_to_text")

        if file_extension == ".pdf":
            # Prefer file_path to avoid reading all bytes into memory
            if file_path:
                with fitz.open(file_path) as pdf_doc:
                    for page in pdf_doc:
                        text += page.get_text()
            else:
                with fitz.open(stream=file_content, filetype="pdf") as pdf_doc:
                    for page in pdf_doc:
                        text += page.get_text()

        elif file_extension == ".docx":
            if file_path:
                doc = Document(file_path)
            else:
                doc = Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif file_extension == ".txt":
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            else:
                # Decode with error handling
                text = file_content.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported file type for resume analysis: {file_extension}")

        if not text.strip():
            raise ValueError("Could not extract any text from the document.")

        return text
    except Exception as e:
        LOG.error(f"Failed to parse document {filename or file_path}: {e}")
        # Re-raise a user-friendly error to be caught by the service layer
        raise ValueError(f"Could not read the content of the file '{filename or file_path}'. It may be corrupted or in an unsupported format.") from e